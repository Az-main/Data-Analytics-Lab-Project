# -*- coding: utf-8 -*-
"""Build the concise 5-page LOCO Anomaly Inspector product report."""

import csv
import os
import shutil
import tempfile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(ROOT, "07_paper_draft")
DASH_DIR = os.path.join(ROOT, "09_dashboard")
EDA_DIR = os.path.join(ROOT, "04_probe_results")
METHOD_DIR = os.path.join(ROOT, "06_method_results")
OUT = os.path.join(REPORT_DIR, "Project_A_LOCO_Report_Revised.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
INK = RGBColor(0x20, 0x25, 0x2B)
MUTED = RGBColor(0x5A, 0x64, 0x70)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E7F4EC"
LIGHT_GREY = "F2F4F7"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(15.5)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(9)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(10.25)
    bullet_style.paragraph_format.left_indent = Inches(0.5)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.line_spacing = 1.08

    number_style = doc.styles["List Number"]
    number_style.font.name = "Calibri"
    number_style.font.size = Pt(10.25)
    number_style.paragraph_format.left_indent = Inches(0.5)
    number_style.paragraph_format.first_line_indent = Inches(-0.25)
    number_style.paragraph_format.space_after = Pt(4)
    number_style.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.text = "LOCO Anomaly Inspector | Final Project Report"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_body(doc, text, after=5):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(after)
    return paragraph


def add_bullet(doc, lead, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    lead_run = paragraph.add_run(lead)
    set_run_font(lead_run, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    return paragraph


def add_step(doc, lead, text):
    paragraph = doc.add_paragraph(style="List Number")
    lead_run = paragraph.add_run(lead)
    set_run_font(lead_run, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.25, italic=True, color=MUTED)
    return paragraph


def add_picture(doc, path, width, caption, alt_text=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    picture = paragraph.add_run().add_picture(path, width=Inches(width))
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", alt_text or caption)
    doc_pr.set("title", caption.split(".", 1)[0])
    paragraph.paragraph_format.keep_with_next = True
    add_caption(doc, caption)


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, bold=True, color=NAVY)
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [3120, 3120, 3120])
    metrics = [
        ("0.761", "Overall AUROC"),
        ("0.731", "Logical AUROC"),
        ("0.804", "Structural AUROC"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_fill(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        value_run = paragraph.add_run(value + "\n")
        set_run_font(value_run, size=17, bold=True, color=NAVY)
        label_run = paragraph.add_run(label)
        set_run_font(label_run, size=8.5, color=MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def read_results():
    path = os.path.join(DASH_DIR, "corrected_main_results.csv")
    with open(path, newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def make_results_chart(rows, path):
    ordered = sorted(rows, key=lambda row: float(row["overall"]), reverse=True)
    canvas = Image.new("RGB", (1500, 610), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(34, bold=True)
    label_font = load_font(23)
    value_font = load_font(22, bold=True)
    axis_font = load_font(19)
    draw.text((750, 30), "Corrected model comparison", fill="#1F3A5F",
              font=title_font, anchor="ma")

    left = 520
    right = 1390
    top = 95
    row_height = 66
    minimum = 0.50
    maximum = 0.80
    for tick in (0.50, 0.60, 0.70, 0.80):
        x = left + int((tick - minimum) / (maximum - minimum) * (right - left))
        draw.line((x, top - 5, x, top + row_height * len(ordered) - 6),
                  fill="#D9DEE5", width=2)
        draw.text((x, 575), f"{tick:.2f}", fill="#5A6470", font=axis_font, anchor="ma")

    for index, row in enumerate(ordered):
        label = row["method"].replace("DINOv2 ", "")
        value = float(row["overall"])
        y = top + index * row_height
        draw.text((500, y + 23), label, fill="#20252B", font=label_font, anchor="ra")
        end = left + int((value - minimum) / (maximum - minimum) * (right - left))
        color = "#2E74B5" if index == 0 else "#8FA3BF"
        draw.rounded_rectangle((left, y + 5, end, y + 42), radius=8, fill=color)
        draw.text((min(end + 12, 1460), y + 23), f"{value:.3f}",
                  fill="#20252B", font=value_font, anchor="lm")

    draw.text(((left + right) // 2, 600), "Overall image-level AUROC",
              fill="#5A6470", font=axis_font, anchor="ms")
    canvas.save(path, quality=95)


def make_pipeline(path):
    canvas = Image.new("RGB", (1620, 350), "white")
    draw = ImageDraw.Draw(canvas)
    label_font = load_font(24, bold=True)
    labels = [
        "Audit and\nletterbox resize",
        "Frozen DINOv2-small\npatch features",
        "Three DINOv2 detectors\n+ one global proxy",
        "Validation-only\nnormalization and fusion",
        "Results, explanations,\nand dashboard",
    ]
    x_positions = [25, 345, 665, 985, 1305]
    for index, (x, label) in enumerate(zip(x_positions, labels)):
        draw.rounded_rectangle((x, 75, x + 270, 270), radius=18,
                               fill="#E8EEF5", outline="#2E74B5", width=4)
        draw.multiline_text((x + 135, 172), label, fill="#20252B",
                            font=label_font, anchor="mm", align="center", spacing=8)
        if index < len(labels) - 1:
            start_x = x + 276
            end_x = x_positions[index + 1] - 8
            draw.line((start_x, 172, end_x, 172), fill="#5A6470", width=5)
            draw.polygon(
                [(end_x, 172), (end_x - 18, 160), (end_x - 18, 184)],
                fill="#5A6470",
            )
    canvas.save(path, quality=95)


def load_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def make_feature_snapshot(path):
    overlay_path = os.path.join(EDA_DIR, "eda_mask_overlay_cleaned.png")
    heatmap_path = os.path.join(
        METHOD_DIR,
        "GridAware_DINOv2",
        "gridaware_anomaly_maps",
        "breakfast_box_structural_anomalies_053.png",
    )
    overlay = Image.open(overlay_path).convert("RGB")
    heatmap = Image.open(heatmap_path).convert("RGB")

    # Use one structural example from the first row of the ground-truth overlay grid.
    cell_width = overlay.width // 5
    overlay_crop = overlay.crop(
        (cell_width * 2, 0, cell_width * 3, int(overlay.height * 0.28))
    )
    overlay_crop.thumbnail((620, 420), Image.Resampling.LANCZOS)
    heatmap = heatmap.resize((390, 390), Image.Resampling.NEAREST)

    canvas = Image.new("RGB", (1240, 500), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(26, bold=True)
    label_font = load_font(21, bold=True)
    note_font = load_font(18)
    draw.text((620, 18), "Examples used for explanation", fill="#1F3A5F",
              font=title_font, anchor="ma")
    canvas.paste(overlay_crop, (120, 75))
    canvas.paste(heatmap, (735, 75))
    draw.text((310, 455), "Ground-truth anomaly overlay", fill="#20252B",
              font=label_font, anchor="mm")
    draw.text((930, 455), "Region-aware model heatmap", fill="#20252B",
              font=label_font, anchor="mm")
    draw.text((620, 488), "The dashboard keeps ground truth and model output clearly separate.",
              fill="#5A6470", font=note_font, anchor="ms")
    canvas.save(path, quality=95)


def add_results_table(doc, rows):
    ordered = sorted(rows, key=lambda row: float(row["overall"]), reverse=True)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [4320, 1260, 1260, 1260, 1260]
    set_table_geometry(table, widths)
    headers = ["Method", "Logical", "Structural", "Overall", "F1"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_GREY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=8.5, bold=True, color=NAVY)

    for index, row in enumerate(ordered):
        cells = table.add_row().cells
        values = [
            row["method"],
            f'{float(row["logical"]):.3f}',
            f'{float(row["structural"]):.3f}',
            f'{float(row["overall"]):.3f}',
            f'{float(row["f1_overall"]):.3f}',
        ]
        for col, (cell, value) in enumerate(zip(cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=8.25, bold=(index == 0))
        if index == 0:
            for cell in cells:
                set_cell_fill(cell, LIGHT_GREEN)
    add_caption(
        doc,
        "Table 1. Corrected image-level results. Mean fusion is the best pre-specified method.",
    )


def add_methods_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [2700, 4380, 2280]
    set_table_geometry(table, widths)
    headers = ["Branch", "What it checks", "Feature source"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_GREY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=8.5, bold=True, color=NAVY)

    rows = [
        ("Patch Memory", "Local patch appearance compared with normal memory.", "DINOv2-small"),
        ("Region-Aware Memory", "Patch appearance plus approximate image position.", "DINOv2-small"),
        ("Composition Histogram", "The global mix of visual words in the image.", "DINOv2-small"),
        ("Global Image-Stat Proxy", "A fast global image-level difference score.", "Non-DINOv2 proxy"),
    ]
    for branch, check, source in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, (branch, check, source)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=8.25)
    add_caption(doc, "Table 2. The four detector branches used in the comparison.")


def build_report():
    rows = read_results()
    temp_dir = tempfile.mkdtemp(prefix="_report_build_", dir=REPORT_DIR)
    results_chart = os.path.join(temp_dir, "results_chart.png")
    pipeline_chart = os.path.join(temp_dir, "pipeline.png")
    feature_snapshot = os.path.join(temp_dir, "feature_snapshot.png")
    make_results_chart(rows, results_chart)
    make_pipeline(pipeline_chart)
    make_feature_snapshot(feature_snapshot)

    try:
        doc = Document()
        configure_document(doc)

        # Page 1: title, introduction, objective.
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(8)
        title.paragraph_format.space_after = Pt(3)
        run = title.add_run("LOCO Anomaly Inspector")
        set_run_font(run, size=23, bold=True, color=NAVY)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(10)
        run = subtitle.add_run(
            "Logical and Structural Anomaly Detection on MVTec LOCO AD"
        )
        set_run_font(run, size=11.5, italic=True, color=MUTED)

        add_metric_strip(doc)

        doc.add_heading("1. Introduction", level=1)
        add_body(
            doc,
            "Factories need to find defective products before they reach customers. "
            "This is difficult because real defects are rare and expensive to label. "
            "The LOCO Anomaly Inspector learns only from normal images and gives each "
            "test image an anomaly score.",
        )
        add_bullet(
            doc,
            "Structural anomalies: ",
            "visible local defects such as damage, contamination, or deformation.",
        )
        add_bullet(
            doc,
            "Logical anomalies: ",
            "missing, extra, misplaced, or incorrectly counted parts. The parts may "
            "look normal by themselves, so these cases are harder.",
        )
        add_body(
            doc,
            "The project uses all five MVTec LOCO AD categories: breakfast box, juice "
            "bottle, pushpins, screw bag, and splicing connectors. The audited dataset "
            "contains 3,651 product images and 1,246 ground-truth masks.",
        )

        doc.add_heading("2. Objective of the Product", level=1)
        add_body(
            doc,
            "The product combines a reproducible machine-learning pipeline with a "
            "Power BI dashboard and a self-contained HTML dashboard. Its main goals are:",
        )
        add_bullet(doc, "Detect defects: ", "rank good and abnormal test images using normal training data.")
        add_bullet(doc, "Explain results: ", "show scores, thresholds, masks, and anomaly heatmaps.")
        add_bullet(doc, "Compare methods: ", "show logical, structural, overall AUROC, and F1.")
        add_bullet(doc, "Support review: ", "present only EDA and model outputs that help explain decisions.")
        add_callout(
            doc,
            "Main result: ",
            "Mean fusion achieved 0.761 overall AUROC. Region-Aware Memory was the "
            "strongest individual detector at 0.750 overall AUROC.",
        )

        doc.add_page_break()

        # Page 2: features and snapshots.
        doc.add_heading("3. Features Description with Relevant Snapshots", level=1)
        add_body(
            doc,
            "The dashboard is available in Power BI and as an offline HTML file. It "
            "brings the dataset summary, model comparison, category results, threshold "
            "metrics, feature selection, EDA, and explanation examples into one place.",
        )
        add_bullet(
            doc,
            "Overview and leaderboard: ",
            "show the best method and compare all seven scoring configurations.",
        )
        add_bullet(
            doc,
            "Category analysis: ",
            "filters results by product category and anomaly type.",
        )
        add_bullet(
            doc,
            "Relevant EDA only: ",
            "keeps sample images, mask overlays, normal mean/variance, and separate "
            "category heatmaps because these views help explain the modelling choices.",
        )
        add_bullet(
            doc,
            "Explainability: ",
            "keeps ground-truth masks separate from model heatmaps and displays the "
            "model score against a validation-derived threshold.",
        )
        add_picture(
            doc,
            feature_snapshot,
            6.35,
            "Figure 1. Relevant explanation views: a ground-truth overlay and a model "
            "anomaly heatmap. They are different objects and are labelled separately.",
            "Two explanation examples: a red ground-truth anomaly overlay on a product "
            "image and a region-aware model heatmap.",
        )
        add_callout(
            doc,
            "Important reading rule: ",
            "The red overlay is the dataset annotation. The heatmap is the model output. "
            "Neither should be described as the other.",
            fill=LIGHT_GREY,
        )

        doc.add_page_break()

        # Page 3: user guide and implementation.
        doc.add_heading("4. User Guidelines", level=1)
        add_step(
            doc,
            "Open the dashboard. ",
            "Use the Power BI file or open 09_dashboard/dashboard.html in a browser.",
        )
        add_step(
            doc,
            "Start with the overview. ",
            "Check the best overall AUROC and the logical-versus-structural difference.",
        )
        add_step(
            doc,
            "Use filters. ",
            "Choose a method, category, and anomaly type to inspect detailed results.",
        )
        add_step(
            doc,
            "Read a decision. ",
            "A score at or above the saved validation threshold is flagged as anomalous.",
        )
        add_step(
            doc,
            "Check the explanation. ",
            "Compare the original image, ground truth, and model heatmap before drawing a conclusion.",
        )

        doc.add_heading("5. Implementation Details", level=1)
        add_picture(
            doc,
            pipeline_chart,
            6.35,
            "Figure 2. The corrected end-to-end workflow.",
            "Five-step workflow from data audit and letterbox resizing through DINOv2 "
            "features, detectors, validation-only fusion, and dashboard outputs.",
        )
        add_body(
            doc,
            "Images are resized to a 384 x 384 canvas with aspect-ratio-preserving "
            "letterboxing. Product images use bilinear interpolation; masks use "
            "nearest-neighbour interpolation and remain binary. Train/good images fit "
            "the normal models. Validation/good scores set normalization and thresholds. "
            "The test split is used only for final evaluation.",
        )
        add_methods_table(doc)
        add_body(
            doc,
            "The three DINOv2 branches use real frozen DINOv2-small patch features. "
            "The Global Image-Stat Detector is a lightweight non-DINOv2 proxy. Scores "
            "are combined with mean, maximum, and rank-average fusion.",
            after=2,
        )

        doc.add_page_break()

        # Page 4: results and feature selection.
        doc.add_heading("5. Implementation Details - Results", level=1)
        add_results_table(doc, rows)
        add_picture(
            doc,
            results_chart,
            6.35,
            "Figure 3. Corrected overall AUROC values from corrected_main_results.csv.",
            "Horizontal bar chart comparing overall AUROC for seven methods. Mean fusion "
            "is highest at 0.761.",
        )
        add_callout(
            doc,
            "Feature selection: ",
            "The Composition Histogram reduced fusion quality, and one Patch Memory "
            "entry was a duplicate. Removing both produced a smaller three-branch "
            "configuration with 0.764 overall AUROC and 0.826 structural AUROC. "
            "Because this analysis used held-out test labels, the pre-specified mean "
            "fusion remains the headline result.",
            fill=LIGHT_GREEN,
        )
        add_body(
            doc,
            "Structural anomalies were easier than logical anomalies. Local damage "
            "changes patch appearance directly, while a wrong number or arrangement of "
            "normal-looking parts requires stronger relational reasoning.",
            after=2,
        )

        doc.add_page_break()

        # Page 5: limitations, conclusion, references.
        doc.add_heading("Current Limitations", level=1)
        add_bullet(
            doc,
            "Single seed: ",
            "the headline experiment has not yet been repeated across several random seeds.",
        )
        add_bullet(
            doc,
            "Letterbox padding: ",
            "padded regions are not fully removed from patch modelling and scoring.",
        )
        add_bullet(
            doc,
            "Logical reasoning: ",
            "nearest-neighbour patch matching does not explicitly count parts or model long-range relationships.",
        )
        add_bullet(
            doc,
            "Localization metrics: ",
            "the final tables do not yet include the complete pixel AUROC and AUPRO evaluation.",
        )
        add_bullet(
            doc,
            "Method scope: ",
            "the global baseline is a proxy, and the patch-memory baseline is PatchCore-style rather than a full official reproduction.",
        )

        doc.add_heading("6. Conclusion", level=1)
        add_body(
            doc,
            "The LOCO Anomaly Inspector meets the main project goal: it trains on "
            "normal images, detects both logical and structural anomalies, compares "
            "multiple ML/DL scoring methods, and explains results through dashboards "
            "and visual evidence.",
        )
        add_body(
            doc,
            "Mean fusion achieved the best pre-specified overall result at 0.761 AUROC. "
            "Region-Aware Memory was the strongest single detector. The work also shows "
            "that keeping every feature is not always helpful: the reduced three-branch "
            "model was slightly better, although it is reported as analysis rather than "
            "the headline model.",
        )
        add_body(
            doc,
            "The next priorities are to mask padded regions, repeat experiments across "
            "multiple seeds, add complete localization metrics, and test stronger "
            "relational methods for counting and arrangement errors.",
        )
        add_callout(
            doc,
            "Final assessment: ",
            "The project delivers a clear, reproducible, and honest anomaly-detection "
            "product. Its strongest contribution is the combination of real DINOv2 "
            "features, leakage-aware evaluation, model comparison, feature selection, "
            "and simple visual explanation.",
        )

        doc.add_heading("References", level=1)
        references = [
            "[1] Bergmann et al. MVTec LOCO AD and GCAD. IJCV, 2022.",
            "[2] Batzner et al. EfficientAD. WACV, 2024.",
            "[3] Damm et al. AnomalyDINO. WACV, 2025.",
            "[4] Hsieh et al. CSAD. BMVC, 2024.",
        ]
        for reference in references:
            paragraph = doc.add_paragraph(reference)
            paragraph.paragraph_format.space_after = Pt(1)
            for run in paragraph.runs:
                set_run_font(run, size=8.25, color=MUTED)

        doc.core_properties.title = "LOCO Anomaly Inspector - Final Project Report"
        doc.core_properties.subject = "Five-page product report"
        doc.core_properties.keywords = "MVTec LOCO AD, DINOv2, anomaly detection"
        doc.save(OUT)
        return OUT
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    print(build_report())
