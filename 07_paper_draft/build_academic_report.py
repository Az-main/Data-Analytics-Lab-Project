# -*- coding: utf-8 -*-
"""Build the five-page academic report and an Overleaf-ready package."""

import csv
import json
import os
import shutil
import zipfile

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_product_report import (
    LIGHT_BLUE,
    LIGHT_GREEN,
    LIGHT_GREY,
    MUTED,
    NAVY,
    TABLE_WIDTH_DXA,
    add_body,
    add_bullet,
    add_callout,
    add_caption,
    add_page_number,
    add_picture,
    load_font,
    read_results,
    set_cell_fill,
    set_run_font,
    set_table_geometry,
)


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(ROOT, "07_paper_draft")
FIGURE_DIR = os.path.join(REPORT_DIR, "figures")
OVERLEAF_DIR = os.path.join(REPORT_DIR, "overleaf")
OVERLEAF_FIGURE_DIR = os.path.join(OVERLEAF_DIR, "figures")
DOCX_OUT = os.path.join(REPORT_DIR, "Project_A_LOCO_Final_Academic_Report.docx")
ZIP_OUT = os.path.join(REPORT_DIR, "LOCO_Overleaf_Report.zip")

BLUE = RGBColor(0x24, 0x5D, 0x91)
INK = RGBColor(0x1E, 0x24, 0x2B)
TEAL = RGBColor(0x0F, 0x76, 0x6E)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing = 1.04

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(5)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(11.2)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(4)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(9.4)
        style.paragraph_format.left_indent = Inches(0.36)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.02

    header = section.header.paragraphs[0]
    header.text = "LOCO Anomaly Inspector | Data Analytics Lab Final Report"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=8, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_rule(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "245D91")
    borders.append(bottom)
    p_pr.append(borders)


def add_compact_bullet(doc, lead, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1.5)
    lead_run = paragraph.add_run(lead)
    set_run_font(lead_run, size=9.4, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=9.4)
    return paragraph


def add_compact_step(doc, lead, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(1.5)
    lead_run = paragraph.add_run(lead)
    set_run_font(lead_run, size=9.3, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=9.3)
    return paragraph


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    metrics = [
        ("0.761", "Overall AUROC"),
        ("0.731", "Logical AUROC"),
        ("0.804", "Structural AUROC"),
        ("5", "Product categories"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_fill(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        value_run = paragraph.add_run(value + "\n")
        set_run_font(value_run, size=15, bold=True, color=NAVY)
        label_run = paragraph.add_run(label)
        set_run_font(label_run, size=7.7, color=MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def make_pipeline(path):
    canvas = Image.new("RGB", (1800, 360), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(32, bold=True)
    label_font = load_font(23, bold=True)
    note_font = load_font(18)
    draw.text(
        (900, 22),
        "End-to-end anomaly-detection workflow",
        fill="#1F3A5F",
        font=title_font,
        anchor="ma",
    )
    labels = [
        ("1. Audit and resize", "384 x 384 letterbox"),
        ("2. Extract features", "frozen DINOv2-small"),
        ("3. Score anomalies", "three DINOv2 branches\n+ one global proxy"),
        ("4. Calibrate and fuse", "validation-good only"),
        ("5. Evaluate and explain", "held-out test + dashboard"),
    ]
    x_positions = [35, 385, 735, 1085, 1435]
    for index, (x, (label, note)) in enumerate(zip(x_positions, labels)):
        draw.rounded_rectangle(
            (x, 88, x + 305, 290),
            radius=18,
            fill="#E8EEF5",
            outline="#245D91",
            width=4,
        )
        draw.text(
            (x + 152, 145),
            label,
            fill="#1F3A5F",
            font=label_font,
            anchor="mm",
        )
        draw.multiline_text(
            (x + 152, 222),
            note,
            fill="#20252B",
            font=note_font,
            anchor="mm",
            align="center",
            spacing=6,
        )
        if index < len(labels) - 1:
            end_x = x_positions[index + 1] - 12
            draw.line((x + 309, 190, end_x, 190), fill="#5A6470", width=5)
            draw.polygon(
                [(end_x, 190), (end_x - 17, 178), (end_x - 17, 202)],
                fill="#5A6470",
            )
    canvas.save(path, quality=95)


def short_method_name(method):
    replacements = {
        "DINOv2 Region-Aware Memory": "Region-Aware Memory",
        "DINOv2 Patch Memory (NN)": "Patch Memory",
        "DINOv2 Composition Histogram (BoVW)": "Composition Histogram",
        "Global Image-Stat Detector (proxy)": "Global Image-Stat proxy",
        "Fusion (rank-avg)": "Fusion (rank average)",
    }
    return replacements.get(method, method)


def make_results_chart(rows, path):
    ordered = sorted(rows, key=lambda row: float(row["overall"]), reverse=True)
    canvas = Image.new("RGB", (1600, 650), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = load_font(34, bold=True)
    label_font = load_font(22)
    value_font = load_font(21, bold=True)
    axis_font = load_font(18)
    draw.text(
        (800, 26),
        "Overall image-level AUROC",
        fill="#1F3A5F",
        font=title_font,
        anchor="ma",
    )

    left, right, top = 525, 1475, 92
    row_height = 68
    minimum, maximum = 0.50, 0.80
    for tick in (0.50, 0.60, 0.70, 0.80):
        x = left + int((tick - minimum) / (maximum - minimum) * (right - left))
        draw.line(
            (x, top - 4, x, top + row_height * len(ordered) - 15),
            fill="#D9DEE5",
            width=2,
        )
        draw.text((x, 620), f"{tick:.2f}", fill="#5A6470", font=axis_font, anchor="ma")

    for index, row in enumerate(ordered):
        label = short_method_name(row["method"])
        value = float(row["overall"])
        y = top + index * row_height
        draw.text((500, y + 22), label, fill="#20252B", font=label_font, anchor="ra")
        end = left + int((value - minimum) / (maximum - minimum) * (right - left))
        color = "#0F766E" if index == 0 else "#7D9CBD"
        draw.rounded_rectangle((left, y + 4, end, y + 42), radius=8, fill=color)
        draw.text(
            (min(end + 12, 1540), y + 22),
            f"{value:.3f}",
            fill="#20252B",
            font=value_font,
            anchor="lm",
        )
    canvas.save(path, quality=95)


def crop_explainability(source, target):
    image = Image.open(source).convert("RGB")
    width, height = image.size
    top = min(250, max(0, height // 4))
    cropped = image.crop((18, top, width - 18, height - 10))
    cropped.save(target, quality=95)


def add_methods_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [2440, 4520, 2400])
    headers = ["Detector branch", "Purpose", "Feature source"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_GREY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=8.1, bold=True, color=NAVY)

    rows = [
        ("Patch Memory", "Find local patches unlike the normal memory bank.", "DINOv2-small"),
        ("Region-Aware Memory", "Compare appearance while keeping rough image position.", "DINOv2-small"),
        ("Composition Histogram", "Measure changes in the global mix of visual words.", "DINOv2-small"),
        ("Global Image-Stat proxy", "Provide a fast whole-image difference score.", "Non-DINOv2 proxy"),
    ]
    for branch, purpose, source in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, (branch, purpose, source)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=7.9)
    add_caption(doc, "Table 1. Detector branches used in the reported comparison.")


def add_results_table(doc, rows):
    ordered = sorted(rows, key=lambda row: float(row["overall"]), reverse=True)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, [4140, 1305, 1305, 1305, 1305])
    headers = ["Method", "Logical", "Structural", "Overall", "F1"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, LIGHT_GREY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=7.9, bold=True, color=NAVY)

    for index, row in enumerate(ordered):
        cells = table.add_row().cells
        values = [
            short_method_name(row["method"]),
            f'{float(row["logical"]):.3f}',
            f'{float(row["structural"]):.3f}',
            f'{float(row["overall"]):.3f}',
            f'{float(row["f1_overall"]):.3f}',
        ]
        for col, (cell, value) in enumerate(zip(cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=7.65, bold=(index == 0))
        if index == 0:
            for cell in cells:
                set_cell_fill(cell, LIGHT_GREEN)
    add_caption(
        doc,
        "Table 2. Corrected image-level results. AUROC is threshold-free; F1 uses a "
        "validation-good operating threshold.",
    )


def add_reference(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Inches(0.13)
    paragraph.paragraph_format.first_line_indent = Inches(-0.13)
    paragraph.paragraph_format.space_after = Pt(1)
    for run in paragraph.runs:
        set_run_font(run, size=7.7, color=MUTED)


def build_docx(rows, pipeline_path, chart_path, explainability_path):
    overview_path = os.path.join(FIGURE_DIR, "dashboard_overview.png")
    doc = Document()
    configure_document(doc)

    # Page 1: title, abstract, introduction, and objective.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("LOCO Anomaly Inspector")
    set_run_font(run, size=23, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run = subtitle.add_run(
        "Logical and Structural Anomaly Detection on MVTec LOCO AD"
    )
    set_run_font(run, size=11.5, italic=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run("Data Analytics Lab | Final Project Report | June 2026")
    set_run_font(run, size=8.8, bold=True, color=TEAL)
    add_rule(doc)

    add_callout(
        doc,
        "Abstract. ",
        "This project detects industrial product defects using only normal training "
        "images. Frozen DINOv2-small features support three anomaly detectors, while a "
        "lightweight global proxy provides an additional baseline. The system compares "
        "seven scoring configurations and presents the results in Power BI and an "
        "offline HTML dashboard. Mean fusion achieved 0.761 overall AUROC, with stronger "
        "performance on structural anomalies than logical anomalies.",
        fill=LIGHT_GREY,
    )

    doc.add_heading("1. Introduction", level=1)
    add_body(
        doc,
        "Industrial inspection systems must find uncommon defects before products reach "
        "customers. In practice, defective examples are limited and expensive to label. "
        "Anomaly detection addresses this problem by learning the appearance of normal "
        "products and assigning a higher score to unusual test images.",
        after=3,
    )
    add_compact_bullet(
        doc,
        "Structural anomalies: ",
        "local visual defects such as damage, contamination, or deformation.",
    )
    add_compact_bullet(
        doc,
        "Logical anomalies: ",
        "missing, extra, misplaced, or incorrectly counted parts that may still look normal.",
    )
    add_body(
        doc,
        "The study uses all five MVTec LOCO AD categories: breakfast box, juice bottle, "
        "pushpins, screw bag, and splicing connectors. The audited data contains 3,651 "
        "images across normal training, normal validation, and held-out test splits.",
        after=3,
    )

    doc.add_heading("2. Objective of the Product", level=1)
    add_body(
        doc,
        "The objective is to build a clear anomaly-detection product that can:",
        after=2,
    )
    add_compact_bullet(doc, "Detect: ", "rank normal and abnormal images without training on defects.")
    add_compact_bullet(doc, "Compare: ", "report logical, structural, overall AUROC, and F1.")
    add_compact_bullet(doc, "Explain: ", "show scores, thresholds, ground truth, and model heatmaps separately.")
    add_compact_bullet(doc, "Support review: ", "place the useful EDA and model evidence in one dashboard.")
    add_metric_strip(doc)

    doc.add_page_break()

    # Page 2: product features and user guidelines.
    doc.add_heading("3. Product Features", level=1)
    add_body(
        doc,
        "The product is available as a Power BI dashboard and as a self-contained HTML "
        "dashboard. Both versions focus on evidence needed to understand the model.",
        after=2,
    )
    add_compact_bullet(
        doc,
        "Overview and leaderboard: ",
        "show the best method and all seven scoring configurations.",
    )
    add_compact_bullet(
        doc,
        "Relevant EDA: ",
        "show split counts, sample images, mask overlays, and category-level patterns.",
    )
    add_compact_bullet(
        doc,
        "Explainability and drill-down: ",
        "show per-image decisions, category results, thresholds, and heatmaps.",
    )
    add_picture(
        doc,
        overview_path,
        5.72,
        "Figure 1. Dashboard overview showing the headline metrics, dataset split, and "
        "method transparency statement.",
        "Dark dashboard overview with AUROC metric cards, category split counts, and a "
        "short explanation of the anomaly-detection workflow.",
    )

    doc.add_heading("4. User Guidelines", level=1)
    add_compact_step(doc, "Open the dashboard. ", "Use the Power BI file or open 09_dashboard/dashboard.html.")
    add_compact_step(doc, "Read the overview. ", "Check overall, logical, and structural AUROC.")
    add_compact_step(doc, "Use a detailed tab. ", "Select the leaderboard, explainability, or per-category view.")
    add_compact_step(doc, "Interpret a decision. ", "A score at or above the saved validation threshold is flagged as anomalous.")

    doc.add_page_break()

    # Page 3: implementation and methodology.
    doc.add_heading("5. Implementation and Methodology", level=1)
    add_picture(
        doc,
        pipeline_path,
        6.55,
        "Figure 2. End-to-end workflow from data preparation to dashboard reporting.",
        "Five-step workflow: audit and resize, extract frozen DINOv2-small features, "
        "score anomalies, calibrate and fuse, and evaluate in the dashboard.",
    )

    doc.add_heading("5.1 Data Preparation", level=2)
    add_body(
        doc,
        "Images are placed on a 384 x 384 canvas using aspect-ratio-preserving "
        "letterboxing. Product images use bilinear interpolation. Ground-truth masks "
        "use nearest-neighbour interpolation and remain binary. Only normal images are "
        "used to fit the anomaly detectors.",
        after=3,
    )

    doc.add_heading("5.2 Model Methods", level=2)
    add_methods_table(doc)
    add_body(
        doc,
        "The three DINOv2 branches use real frozen DINOv2-small patch features. The "
        "Global Image-Stat detector is a lightweight non-DINOv2 proxy. Mean, maximum, "
        "and rank-average fusion combine the branch scores. This report does not claim "
        "an official EfficientAD or PatchCore implementation.",
        after=3,
    )

    doc.add_heading("5.3 Evaluation Protocol", level=2)
    add_compact_bullet(doc, "Training: ", "fit normal memories using train/good images.")
    add_compact_bullet(doc, "Calibration: ", "use validation/good scores for normalization and the 90th-percentile threshold.")
    add_compact_bullet(doc, "Testing: ", "use the held-out test split only for the final reported evaluation.")
    add_compact_bullet(doc, "Metrics: ", "report image-level AUROC and F1; AUROC is the main comparison metric.")

    doc.add_page_break()

    # Page 4: results and feature selection.
    doc.add_heading("6. Results and Feature Selection", level=1)
    add_results_table(doc, rows)
    add_picture(
        doc,
        chart_path,
        6.35,
        "Figure 3. Corrected overall AUROC for all reported methods and fusion rules.",
        "Horizontal bar chart comparing overall image-level AUROC. Mean fusion is "
        "highest at 0.761.",
    )
    add_compact_bullet(
        doc,
        "Best pre-specified result: ",
        "mean fusion reached 0.761 overall AUROC.",
    )
    add_compact_bullet(
        doc,
        "Best single detector: ",
        "Region-Aware Memory reached 0.750 overall AUROC.",
    )
    add_compact_bullet(
        doc,
        "Main pattern: ",
        "structural AUROC (0.804) was higher than logical AUROC (0.731).",
    )
    add_callout(
        doc,
        "Feature-selection analysis. ",
        "Removing a duplicate Patch Memory entry and the weak Composition Histogram "
        "branch produced a three-branch score of 0.764 overall AUROC and 0.826 "
        "structural AUROC. This analysis used held-out test labels, so it is reported "
        "as a later diagnostic result. Mean fusion remains the headline result.",
        fill=LIGHT_GREEN,
    )

    doc.add_page_break()

    # Page 5: explainability, limitations, conclusion, and references.
    doc.add_heading("7. Explainability and Limitations", level=1)
    add_picture(
        doc,
        explainability_path,
        5.85,
        "Figure 4. Per-image dashboard decisions with the anomaly score, validation "
        "threshold, predicted class, and ground-truth defect region.",
        "Dashboard cards showing normal, logical-anomaly, and structural-anomaly "
        "examples with scores and validation thresholds.",
    )
    add_body(
        doc,
        "The red region is the dataset annotation, while the score and prediction are "
        "model outputs. Keeping them separate prevents the ground truth from being "
        "mistaken for a generated explanation.",
        after=2,
    )
    add_compact_bullet(doc, "Single run: ", "the headline experiment has not been repeated across several random seeds.")
    add_compact_bullet(doc, "Padding: ", "letterbox padding is not fully removed from patch modelling and scoring.")
    add_compact_bullet(doc, "Logical reasoning: ", "nearest-neighbour features do not explicitly count parts or model long-range relations.")
    add_compact_bullet(doc, "Localization: ", "the final comparison does not yet include complete pixel AUROC and AUPRO results.")

    doc.add_heading("8. Conclusion", level=1)
    add_body(
        doc,
        "The LOCO Anomaly Inspector meets the project requirements by combining "
        "unsupervised anomaly detection, method comparison, useful EDA, and an "
        "interactive dashboard. Mean fusion achieved 0.761 overall AUROC, and the "
        "Region-Aware branch was the strongest single detector. The results also show "
        "that logical anomalies remain harder than visible structural defects.",
        after=2,
    )

    doc.add_heading("References", level=1)
    add_reference(doc, "[1] Bergmann et al., MVTec LOCO AD and GCAD, International Journal of Computer Vision, 2022.")
    add_reference(doc, "[2] Oquab et al., DINOv2: Learning Robust Visual Features without Supervision, TMLR, 2024.")
    add_reference(doc, "[3] MVTec LOCO AD dataset documentation and project experiment outputs.")

    doc.core_properties.title = "LOCO Anomaly Inspector - Final Academic Report"
    doc.core_properties.subject = "Five-page Data Analytics Lab final project report"
    doc.core_properties.keywords = "MVTec LOCO AD, DINOv2, anomaly detection, dashboard"
    doc.save(DOCX_OUT)


def latex_escape(text):
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
    }
    return "".join(replacements.get(char, char) for char in text)


def latex_result_rows(rows):
    ordered = sorted(rows, key=lambda row: float(row["overall"]), reverse=True)
    lines = []
    for index, row in enumerate(ordered):
        values = [
            latex_escape(short_method_name(row["method"])),
            f'{float(row["logical"]):.3f}',
            f'{float(row["structural"]):.3f}',
            f'{float(row["overall"]):.3f}',
            f'{float(row["f1_overall"]):.3f}',
        ]
        line = " & ".join(values) + r" \\"
        if index == 0:
            line = r"\rowcolor{bestgreen} " + line
        lines.append(line)
    return "\n".join(lines)


def build_latex(rows):
    result_rows = latex_result_rows(rows)
    tex = r"""\documentclass[10pt]{article}
\usepackage[letterpaper,margin=0.68in,headheight=14pt]{geometry}
\usepackage[T1]{fontenc}
\usepackage{lmodern}
\usepackage{microtype}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{tabularx}
\usepackage{array}
\usepackage[table]{xcolor}
\usepackage{enumitem}
\usepackage{caption}
\usepackage{float}
\usepackage{titlesec}
\usepackage{fancyhdr}
\usepackage[hidelinks]{hyperref}

\definecolor{navy}{HTML}{1F3A5F}
\definecolor{blue}{HTML}{245D91}
\definecolor{teal}{HTML}{0F766E}
\definecolor{lightblue}{HTML}{E8EEF5}
\definecolor{lightgrey}{HTML}{F2F4F7}
\definecolor{bestgreen}{HTML}{E7F4EC}

\pagestyle{fancy}
\fancyhf{}
\lhead{\footnotesize LOCO Anomaly Inspector | Data Analytics Lab Final Report}
\rfoot{\footnotesize Page \thepage}
\renewcommand{\headrulewidth}{0.3pt}
\renewcommand{\footrulewidth}{0pt}

\titleformat{\section}{\large\bfseries\color{blue}}{\thesection.}{0.45em}{}
\titleformat{\subsection}{\normalsize\bfseries\color{navy}}{\thesubsection}{0.45em}{}
\titlespacing*{\section}{0pt}{5pt}{3pt}
\titlespacing*{\subsection}{0pt}{4pt}{2pt}
\setlist[itemize]{leftmargin=1.25em,itemsep=1pt,topsep=2pt}
\setlist[enumerate]{leftmargin=1.35em,itemsep=1pt,topsep=2pt}
\captionsetup{font=footnotesize,labelfont=bf,skip=2pt}
\setlength{\parindent}{0pt}
\setlength{\parskip}{3pt}
\renewcommand{\arraystretch}{1.08}

\begin{document}
\fontsize{9.4}{10.8}\selectfont

% Add student name, ID, and institution below if your faculty requires them.
\begin{center}
{\LARGE\bfseries\color{navy} LOCO Anomaly Inspector}\\[2pt]
{\large\itshape Logical and Structural Anomaly Detection on MVTec LOCO AD}\\[3pt]
{\small\bfseries\color{teal} Data Analytics Lab | Final Project Report | June 2026}
\end{center}
\vspace{-2pt}
\hrule height 0.8pt
\vspace{5pt}

\colorbox{lightgrey}{%
\parbox{\dimexpr\linewidth-2\fboxsep\relax}{%
\textbf{Abstract.} This project detects industrial product defects using only normal
training images. Frozen DINOv2-small features support three anomaly detectors, while a
lightweight global proxy provides an additional baseline. The system compares seven
scoring configurations and presents the results in Power BI and an offline HTML
dashboard. Mean fusion achieved 0.761 overall AUROC, with stronger performance on
structural anomalies than logical anomalies.}}

\section{Introduction}
Industrial inspection systems must find uncommon defects before products reach
customers. In practice, defective examples are limited and expensive to label.
Anomaly detection addresses this problem by learning the appearance of normal products
and assigning a higher score to unusual test images.

\begin{itemize}
\item \textbf{Structural anomalies:} local visual defects such as damage,
contamination, or deformation.
\item \textbf{Logical anomalies:} missing, extra, misplaced, or incorrectly counted
parts that may still look normal.
\end{itemize}

The study uses all five MVTec LOCO AD categories: breakfast box, juice bottle,
pushpins, screw bag, and splicing connectors. The audited data contains 3,651 images
across normal training, normal validation, and held-out test splits.

\section{Objective of the Product}
The objective is to build a clear anomaly-detection product that can:
\begin{itemize}
\item \textbf{Detect} normal and abnormal images without training on defects.
\item \textbf{Compare} logical, structural, overall AUROC, and F1.
\item \textbf{Explain} scores, thresholds, ground truth, and model heatmaps separately.
\item \textbf{Support review} by placing useful EDA and model evidence in one dashboard.
\end{itemize}

\begin{center}
\rowcolors{1}{lightblue}{lightblue}
\begin{tabular}{>{\centering\arraybackslash}p{0.22\linewidth}
                >{\centering\arraybackslash}p{0.22\linewidth}
                >{\centering\arraybackslash}p{0.22\linewidth}
                >{\centering\arraybackslash}p{0.22\linewidth}}
\textbf{\large 0.761} & \textbf{\large 0.731} & \textbf{\large 0.804} & \textbf{\large 5}\\
\footnotesize Overall AUROC & \footnotesize Logical AUROC &
\footnotesize Structural AUROC & \footnotesize Product categories
\end{tabular}
\end{center}

\newpage
\section{Product Features}
The product is available as a Power BI dashboard and as a self-contained HTML
dashboard. Both versions focus on evidence needed to understand the model.
\begin{itemize}
\item \textbf{Overview and leaderboard:} show the best method and all seven scoring configurations.
\item \textbf{Relevant EDA:} show split counts, sample images, mask overlays, and category patterns.
\item \textbf{Explainability and drill-down:} show per-image decisions, category results, thresholds, and heatmaps.
\end{itemize}

\begin{figure}[H]
\centering
\includegraphics[width=0.86\linewidth]{figures/dashboard_overview.png}
\caption{Dashboard overview showing the headline metrics, dataset split, and method transparency statement.}
\end{figure}

\section{User Guidelines}
\begin{enumerate}
\item \textbf{Open the dashboard.} Use the Power BI file or open
\texttt{09\_dashboard/dashboard.html}.
\item \textbf{Read the overview.} Check overall, logical, and structural AUROC.
\item \textbf{Use a detailed tab.} Select the leaderboard, explainability, or per-category view.
\item \textbf{Interpret a decision.} A score at or above the saved validation threshold is flagged as anomalous.
\end{enumerate}

\newpage
\section{Implementation and Methodology}
\begin{figure}[H]
\centering
\includegraphics[width=\linewidth]{figures/pipeline.png}
\caption{End-to-end workflow from data preparation to dashboard reporting.}
\end{figure}

\subsection{Data Preparation}
Images are placed on a 384 $\times$ 384 canvas using aspect-ratio-preserving
letterboxing. Product images use bilinear interpolation. Ground-truth masks use
nearest-neighbour interpolation and remain binary. Only normal images are used to fit
the anomaly detectors.

\subsection{Model Methods}
\begin{table}[H]
\centering
\footnotesize
\begin{tabularx}{\linewidth}{>{\bfseries}p{0.23\linewidth}X p{0.20\linewidth}}
\toprule
Detector branch & Purpose & Feature source\\
\midrule
Patch Memory & Find local patches unlike the normal memory bank. & DINOv2-small\\
Region-Aware Memory & Compare appearance while keeping rough image position. & DINOv2-small\\
Composition Histogram & Measure changes in the global mix of visual words. & DINOv2-small\\
Global Image-Stat proxy & Provide a fast whole-image difference score. & Non-DINOv2 proxy\\
\bottomrule
\end{tabularx}
\caption{Detector branches used in the reported comparison.}
\end{table}

The three DINOv2 branches use real frozen DINOv2-small patch features. The Global
Image-Stat detector is a lightweight non-DINOv2 proxy. Mean, maximum, and rank-average
fusion combine the branch scores. This report does not claim an official EfficientAD
or PatchCore implementation.

\subsection{Evaluation Protocol}
\begin{itemize}
\item \textbf{Training:} fit normal memories using train/good images.
\item \textbf{Calibration:} use validation/good scores for normalization and the 90th-percentile threshold.
\item \textbf{Testing:} use the held-out test split only for final evaluation.
\item \textbf{Metrics:} report image-level AUROC and F1; AUROC is the main comparison metric.
\end{itemize}

\newpage
\section{Results and Feature Selection}
\begin{table}[H]
\centering
\scriptsize
\begin{tabularx}{\linewidth}{Xrrrr}
\toprule
Method & Logical & Structural & Overall & F1\\
\midrule
""" + result_rows + r"""
\bottomrule
\end{tabularx}
\caption{Corrected image-level results. AUROC is threshold-free; F1 uses a validation-good operating threshold.}
\end{table}

\begin{figure}[H]
\centering
\includegraphics[width=0.95\linewidth]{figures/results_chart.png}
\caption{Corrected overall AUROC for all reported methods and fusion rules.}
\end{figure}

\begin{itemize}
\item \textbf{Best pre-specified result:} mean fusion reached 0.761 overall AUROC.
\item \textbf{Best single detector:} Region-Aware Memory reached 0.750 overall AUROC.
\item \textbf{Main pattern:} structural AUROC (0.804) was higher than logical AUROC (0.731).
\end{itemize}

\colorbox{bestgreen}{%
\parbox{\dimexpr\linewidth-2\fboxsep\relax}{%
\textbf{Feature-selection analysis.} Removing a duplicate Patch Memory entry and the
weak Composition Histogram branch produced a three-branch score of 0.764 overall AUROC
and 0.826 structural AUROC. This analysis used held-out test labels, so it is reported
as a later diagnostic result. Mean fusion remains the headline result.}}

\newpage
\section{Explainability and Limitations}
\begin{figure}[H]
\centering
\includegraphics[width=0.88\linewidth]{figures/explainability_decisions.png}
\caption{Per-image decisions with anomaly score, validation threshold, predicted class, and ground-truth defect region.}
\end{figure}

The red region is the dataset annotation, while the score and prediction are model
outputs. Keeping them separate prevents the ground truth from being mistaken for a
generated explanation.

\begin{itemize}
\item \textbf{Single run:} the headline experiment has not been repeated across several random seeds.
\item \textbf{Padding:} letterbox padding is not fully removed from patch modelling and scoring.
\item \textbf{Logical reasoning:} nearest-neighbour features do not explicitly count parts or model long-range relations.
\item \textbf{Localization:} the final comparison does not yet include complete pixel AUROC and AUPRO results.
\end{itemize}

\section{Conclusion}
The LOCO Anomaly Inspector meets the project requirements by combining unsupervised
anomaly detection, method comparison, useful EDA, and an interactive dashboard. Mean
fusion achieved 0.761 overall AUROC, and the Region-Aware branch was the strongest
single detector. The results also show that logical anomalies remain harder than
visible structural defects.

\section*{References}
\footnotesize
[1] Bergmann et al., MVTec LOCO AD and GCAD, \textit{International Journal of Computer Vision}, 2022.\\
[2] Oquab et al., DINOv2: Learning Robust Visual Features without Supervision, \textit{TMLR}, 2024.\\
[3] MVTec LOCO AD dataset documentation and project experiment outputs.

\end{document}
"""
    os.makedirs(OVERLEAF_DIR, exist_ok=True)
    with open(os.path.join(OVERLEAF_DIR, "main.tex"), "w", encoding="utf-8") as handle:
        handle.write(tex)

    readme = """LOCO Anomaly Inspector - Overleaf Package

1. Upload the full contents of this folder to a new Overleaf project.
2. Keep main.tex as the main document.
3. Add your name, student ID, course code, and institution below the title if required.
4. Compile with pdfLaTeX.

The report is arranged as five explicit pages. If Overleaf moves a small amount of
text to a sixth page, reduce one large figure from 0.88/0.95 linewidth to 0.84/0.90,
or change the document font size from 9.4 to 9.2 in main.tex.

The figures and numerical results were copied from the local project evidence.
"""
    with open(os.path.join(OVERLEAF_DIR, "README.txt"), "w", encoding="utf-8") as handle:
        handle.write(readme)


def package_overleaf(pipeline_path, chart_path, explainability_path):
    os.makedirs(OVERLEAF_FIGURE_DIR, exist_ok=True)
    figure_sources = {
        os.path.join(FIGURE_DIR, "dashboard_overview.png"): "dashboard_overview.png",
        pipeline_path: "pipeline.png",
        chart_path: "results_chart.png",
        explainability_path: "explainability_decisions.png",
    }
    for source, name in figure_sources.items():
        shutil.copy2(source, os.path.join(OVERLEAF_FIGURE_DIR, name))

    with zipfile.ZipFile(ZIP_OUT, "w", zipfile.ZIP_DEFLATED) as archive:
        for base, _, files in os.walk(OVERLEAF_DIR):
            for filename in files:
                path = os.path.join(base, filename)
                archive.write(path, os.path.relpath(path, OVERLEAF_DIR))


def structural_qa():
    document = Document(DOCX_OUT)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading")
    ]
    page_breaks = sum(
        1
        for paragraph in document.paragraphs
        for node in paragraph._p.xpath(".//w:br[@w:type='page']")
    )
    inline_shapes = len(document.inline_shapes)
    tables = len(document.tables)
    required_headings = [
        "1. Introduction",
        "2. Objective of the Product",
        "3. Product Features",
        "4. User Guidelines",
        "5. Implementation and Methodology",
        "6. Results and Feature Selection",
        "7. Explainability and Limitations",
        "8. Conclusion",
        "References",
    ]
    missing = [heading for heading in required_headings if heading not in headings]

    main_tex = os.path.join(OVERLEAF_DIR, "main.tex")
    with open(main_tex, encoding="utf-8") as handle:
        tex = handle.read()
    tex_assets = [
        "dashboard_overview.png",
        "pipeline.png",
        "results_chart.png",
        "explainability_decisions.png",
    ]
    missing_assets = [
        name
        for name in tex_assets
        if not os.path.exists(os.path.join(OVERLEAF_FIGURE_DIR, name))
    ]
    brace_balance = tex.count("{") - tex.count("}")
    report = {
        "docx": DOCX_OUT,
        "docx_page_breaks": page_breaks,
        "docx_expected_pages": page_breaks + 1,
        "docx_headings": headings,
        "docx_missing_headings": missing,
        "docx_inline_images": inline_shapes,
        "docx_tables": tables,
        "latex": main_tex,
        "latex_newpage_count": tex.count(r"\newpage"),
        "latex_brace_balance": brace_balance,
        "latex_missing_assets": missing_assets,
        "zip": ZIP_OUT,
    }
    qa_path = os.path.join(REPORT_DIR, "academic_report_qa.json")
    with open(qa_path, "w", encoding="utf-8") as handle:
        json.dump(report, handle, indent=2)
    return report


def build():
    rows = read_results()
    os.makedirs(OVERLEAF_FIGURE_DIR, exist_ok=True)
    pipeline_path = os.path.join(FIGURE_DIR, "academic_pipeline.png")
    chart_path = os.path.join(FIGURE_DIR, "academic_results_chart.png")
    explainability_path = os.path.join(FIGURE_DIR, "explainability_decisions.png")
    make_pipeline(pipeline_path)
    make_results_chart(rows, chart_path)
    crop_explainability(
        os.path.join(FIGURE_DIR, "dashboard_explainability.png"),
        explainability_path,
    )
    build_docx(rows, pipeline_path, chart_path, explainability_path)
    build_latex(rows)
    package_overleaf(pipeline_path, chart_path, explainability_path)
    return structural_qa()


if __name__ == "__main__":
    print(json.dumps(build(), indent=2))
